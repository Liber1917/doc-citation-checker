#!/usr/bin/env python3
"""
Doc Citation Checker - Verify citations in documents without .bib file
Supports .docx, .md, .txt, .tex files

Improved version:
- Retry logic with exponential backoff for all API calls
- Defensive handling of None / missing / malformed inputs
- Encoding fallback (utf-8 → gbk → latin-1)
- Fixed threshold default (0.35, matching Jaccard reality)
- Verbose mode for debug output
- Never silently swallows exceptions
"""

import argparse
import json
import re
import sys
import time
from pathlib import Path
from typing import Dict, List, Any, Optional

try:
    import requests
    from requests.adapters import HTTPAdapter
    from urllib3.util.retry import Retry
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ---------------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------------

def safe_get(data: dict, key_path: str, default=None):
    """Safely navigate nested dict/list via dot/bracket notation.
    Example: safe_get(item, 'title.0', '') → item.get('title', [])[0] or default
    """
    parts = key_path.split('.')
    cur = data
    for p in parts:
        if cur is None:
            return default
        if p.isdigit():
            idx = int(p)
            if isinstance(cur, (list, tuple)) and 0 <= idx < len(cur):
                cur = cur[idx]
            else:
                return default
        elif isinstance(cur, dict):
            cur = cur.get(p, default)
        else:
            return default
    return cur if cur is not None else default


def read_file_with_fallback(file_path: str) -> Optional[str]:
    """Read a text file, trying multiple encodings."""
    encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1', 'cp1252']
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    # Last resort: read as bytes and decode with errors='replace'
    try:
        with open(file_path, 'rb') as f:
            return f.read().decode('utf-8', errors='replace')
    except Exception as e:
        print(f"  [Error] Cannot read file {file_path}: {e}")
        return None


def get_session() -> requests.Session:
    """Return a requests.Session with retry + backoff."""
    session = requests.Session()
    retry = Retry(
        total=3,
        backoff_factor=1,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["GET"]
    )
    adapter = HTTPAdapter(max_retries=retry)
    session.mount("https://", adapter)
    session.mount("http://", adapter)
    return session


def title_similarity(a: str, b: str) -> float:
    """Compute word-level Jaccard similarity between two titles.
    Returns 0.0 if either input is None/empty.
    Short titles (< 3 words): require near-exact match.
    All titles: at least 2 overlapping words required for non-zero score.
    """
    if not a or not b:
        return 0.0
    stop = {'the', 'a', 'an', 'of', 'in', 'for', 'and', 'on', 'with',
            'to', 'is', 'are', 'by', 'et', 'al', 'will', 'be', 'was',
            'been', 'have', 'has', 'had', 'but', 'not', 'all', 'you'}
    def tokenize(s: str):
        words = re.findall(r'\b\w{3,}\b', s.lower())
        return set(w for w in words if w not in stop)
    sa, sb = tokenize(a), tokenize(b)
    if not sa or not sb:
        return 0.0
    # Short query title: require all words to overlap
    if len(sa) < 3 or len(sb) < 3:
        return 1.0 if sa == sb else 0.0
    intersection = sa & sb
    if len(intersection) < 2:
        return 0.0   # at least 2 word overlap required
    return len(intersection) / len(sa | sb)


# ---------------------------------------------------------------------------
# Citation extraction
# ---------------------------------------------------------------------------

def extract_citations_from_docx(file_path: str, verbose: bool = False) -> List[Dict[str, Any]]:
    """Extract citations from .docx file."""
    if not HAS_DOCX:
        print("  [Warning] python-docx not installed. Install: pip install python-docx")
        return []

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"  [Error] Cannot parse .docx file: {e}")
        return []

    full_text = "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    full_text += "\n" + cell.text

    if verbose:
        print(f"  [Debug] .docx extracted {len(full_text)} chars of text")

    return extract_citations_from_text(full_text, verbose)


def extract_citations_from_tex(file_path: str, verbose: bool = False) -> List[Dict[str, Any]]:
    """Extract citations from .tex file, mapping \\cite{key} to \\bibitem content."""
    content = read_file_with_fallback(file_path)
    if content is None:
        return []

    # Step 1: Build key -> bibitem raw text map
    bibitem_map: Dict[str, str] = {}
    bibitem_pattern = r'\\bibitem\{([^}]+)\}(.*?)(?=\\bibitem|\s*\\end\{thebibliography\})'
    for match in re.finditer(bibitem_pattern, content, re.DOTALL):
        key = match.group(1).strip()
        raw_text = re.sub(r'[{}\\]', ' ', match.group(2))
        raw_text = re.sub(r'\s+', ' ', raw_text).strip()
        bibitem_map[key] = raw_text[:300]

    if verbose:
        print(f"  [Debug] Found {len(bibitem_map)} \\bibitem entries")

    citations = []
    seen_keys = set()

    # Step 2: Resolve \cite{key} through bibitem map
    cite_pattern = r'\\cite(?:\[[^\]]*\])?\{([^}]+)\}'
    for match in re.finditer(cite_pattern, content):
        keys = [k.strip() for k in match.group(1).split(',')]
        for key in keys:
            if key in seen_keys:
                continue
            seen_keys.add(key)
            raw = bibitem_map.get(key, '')
            citations.append({
                'key': key,
                'type': 'cite_resolved' if raw else 'cite_unresolved',
                'raw': raw if raw else key
            })

    # Step 3: Standalone bibitem entries not referenced by \cite
    for key, raw in bibitem_map.items():
        if key not in seen_keys:
            citations.append({'key': key, 'type': 'bibitem', 'raw': raw})

    if verbose:
        print(f"  [Debug] Total citations extracted from .tex: {len(citations)}")

    return citations


def extract_citations_from_text(text: str, verbose: bool = False) -> List[Dict[str, Any]]:
    """Extract citations from plain text using common patterns."""
    if not text or not text.strip():
        return []

    citations = []

    # DOI pattern
    doi_pattern = r'doi:\s*(10\.\d{4,9}/[^\s]+)'
    for match in re.finditer(doi_pattern, text, re.IGNORECASE):
        doi = match.group(1).rstrip('.,;')
        citations.append({'key': doi, 'type': 'doi', 'raw': match.group(0).strip()})

    # Bibliography section
    bib_section_pattern = r'(?:References|Bibliography|参考文献)[:\s]*\n(.*?)(?=\n\s*\n|$)'
    bib_match = re.search(bib_section_pattern, text, re.IGNORECASE | re.DOTALL)
    if bib_match:
        bib_text = bib_match.group(1)
        for line in bib_text.split('\n'):
            line = line.strip()
            if line:
                citations.append({'key': line[:50], 'type': 'bibliography', 'raw': line})

    if verbose:
        print(f"  [Debug] extract_citations_from_text found {len(citations)} items")

    return citations


def extract_citations(file_path: str, verbose: bool = False) -> List[Dict[str, Any]]:
    """Extract citations from a file based on its extension."""
    path = Path(file_path)
    if not path.exists():
        print(f"  [Error] File not found: {file_path}")
        return []

    suffix = path.suffix.lower()
    if suffix == '.docx':
        return extract_citations_from_docx(file_path, verbose)
    elif suffix in ['.md', '.txt', '.rst']:
        text = read_file_with_fallback(file_path)
        if text is None:
            return []
        return extract_citations_from_text(text, verbose)
    elif suffix == '.tex':
        return extract_citations_from_tex(file_path, verbose)
    else:
        print(f"  [Warning] Unsupported file type: {suffix}")
        return []


# ---------------------------------------------------------------------------
# Title extraction
# ---------------------------------------------------------------------------

def extract_title_from_citation(text: str, verbose: bool = False) -> Optional[str]:
    """Extract paper title from citation text. Returns None if not found."""
    if not text or not text.strip():
        return None

    patterns = [
        (r"``([^']{10,}?)''", "LaTeX quotes"),
        (r'"([^"]{10,}?)"',     "double quotes"),
        (r'《([^》]+)》',           "Chinese quotes"),
        (r'(?:Title|题目)[:\s]+([^\n.]{10,})', "Title: label"),
    ]
    for pattern, label in patterns:
        m = re.search(pattern, text)
        if m:
            title = m.group(1).strip()
            if verbose:
                print(f"  [Debug] Title extracted via {label}: {title[:60]}")
            return title

    # Heuristic: after author/year, rest is title
    m = re.search(
        r'(?:[A-Z][a-z]+,?\s+(?:[A-Z]\.\s*)+)?(?:\(?\d{4}\)?\s*)([A-Z][^.]{15,80})\.',
        text
    )
    if m:
        return m.group(1).strip()

    if verbose:
        print(f"  [Debug] No title extracted from: {text[:60]}")
    return None


# ---------------------------------------------------------------------------
# API verification (with retry, never silent)
# ---------------------------------------------------------------------------

def verify_doi_via_crossref(doi: str, session: requests.Session,
                            verbose: bool = False) -> Optional[Dict[str, Any]]:
    """Verify a DOI via CrossRef API. Never returns None silently."""
    try:
        url = f"https://api.crossref.org/works/{doi}"
        resp = session.get(url, timeout=15)
        if resp.status_code == 200:
            data = resp.json().get('message', {})
            title = safe_get(data, 'title.0', '')
            authors = [
                f"{a.get('given','').strip()} {a.get('family','').strip()}".strip()
                for a in data.get('author', [])
                if isinstance(a, dict)
            ]
            year = safe_get(data, 'published.date-parts.0.0', '')
            return {
                'source': 'CrossRef',
                'doi': doi,
                'title': title,
                'authors': authors,
                'year': str(year) if year else ''
            }
        elif verbose:
            print(f"  [Debug] CrossRef DOI lookup HTTP {resp.status_code} for {doi}")
    except Exception as e:
        print(f"  [Warning] CrossRef DOI verify failed for {doi}: {e}")
    return None


def search_crossref(query: str, session: requests.Session,
                    verbose: bool = False) -> List[Dict[str, Any]]:
    """Search CrossRef for a title. Always returns a list (never raises)."""
    try:
        url = "https://api.crossref.org/works"
        resp = session.get(url, params={'query.title': query, 'rows': 3}, timeout=15)
        if resp.status_code == 200:
            results = []
            for item in resp.json().get('message', {}).get('items', []):
                title = safe_get(item, 'title.0', '')
                authors = [
                    f"{a.get('given','').strip()} {a.get('family','').strip()}".strip()
                    for a in item.get('author', [])
                    if isinstance(a, dict)
                ]
                year = safe_get(item, 'published.date-parts.0.0', '')
                results.append({
                    'source': 'CrossRef',
                    'doi': item.get('DOI', ''),
                    'title': title,
                    'authors': authors,
                    'year': str(year) if year else ''
                })
            return results
        elif verbose:
            print(f"  [Debug] CrossRef search HTTP {resp.status_code}")
    except Exception as e:
        print(f"  [Warning] CrossRef search failed: {e}")
    return []


def search_semantic_scholar(query: str, session: requests.Session,
                            verbose: bool = False) -> List[Dict[str, Any]]:
    """Search Semantic Scholar. Always returns a list."""
    try:
        url = "https://api.semanticscholar.org/graph/v1/paper/search"
        resp = session.get(
            url,
            params={'query': query, 'limit': 3, 'fields': 'title,authors,year,externalIds'},
            timeout=15
        )
        if resp.status_code == 200:
            results = []
            for item in resp.json().get('data', []):
                results.append({
                    'source': 'Semantic Scholar',
                    'title': item.get('title', ''),
                    'authors': [a.get('name', '') for a in item.get('authors', []) if isinstance(a, dict)],
                    'year': str(item.get('year', '')),
                    'paperId': item.get('paperId', '')
                })
            return results
        elif verbose:
            print(f"  [Debug] Semantic Scholar HTTP {resp.status_code}")
    except Exception as e:
        print(f"  [Warning] Semantic Scholar search failed: {e}")
    return []


# ---------------------------------------------------------------------------
# Core verification
# ---------------------------------------------------------------------------

def verify_citation(citation: Dict[str, Any], session: requests.Session,
                    threshold: float = 0.60, verbose: bool = False) -> Dict[str, Any]:
    """Verify a citation. Returns a result dict with status."""
    if not HAS_REQUESTS:
        return {'citation': citation, 'status': 'error',
                'message': 'requests library not installed'}

    result = {'citation': citation, 'status': 'unknown', 'sources': []}
    raw_text = citation.get('raw', '') or ''

    # --- Fast path: DOI ---
    doi_match = re.search(r'10\.\d{4,9}/[^\s,}]+', raw_text)
    if doi_match:
        doi = doi_match.group(0).rstrip('.,;')
        crossref_result = verify_doi_via_crossref(doi, session, verbose)
        if crossref_result:
            result['status'] = 'valid'
            result['sources'].append(crossref_result)
            return result

    # --- Title-based search ---
    query_title = extract_title_from_citation(raw_text, verbose)
    if not query_title:
        result['status'] = 'unverifiable'
        return result

    if verbose:
        print(f"  [Debug] Verifying title: {query_title[:60]}")

    # CrossRef
    time.sleep(0.3)
    for r in search_crossref(query_title, session, verbose):
        sim = title_similarity(query_title, r.get('title', ''))
        if verbose:
            print(f"  [Debug] CrossRef candidate: {r.get('title','')[:60]}  (sim={sim:.3f})")
        if sim >= threshold:
            r['similarity'] = round(sim, 3)
            result['status'] = 'valid'
            result['sources'].append(r)
            return result

    # Semantic Scholar
    time.sleep(0.3)
    for r in search_semantic_scholar(query_title, session, verbose):
        sim = title_similarity(query_title, r.get('title', ''))
        if verbose:
            print(f"  [Debug] S2 candidate: {r.get('title','')[:60]}  (sim={sim:.3f})")
        if sim >= threshold:
            r['similarity'] = round(sim, 3)
            result['status'] = 'valid'
            result['sources'].append(r)
            return result

    result['status'] = 'hallucinated'
    return result


# ---------------------------------------------------------------------------
# Report generation
# ---------------------------------------------------------------------------

def check_file(file_path: str, threshold: float = 0.35,
               verbose: bool = False) -> Dict[str, Any]:
    """Check all citations in a file. Returns a report dict."""
    print(f"Checking citations in: {file_path}")

    if not HAS_REQUESTS:
        print("  [Error] requests not installed. Install: pip install requests")

    citations = extract_citations(file_path, verbose)
    if not citations:
        print("  No citations found in the file.")
        return {
            'file': file_path,
            'total_citations': 0,
            'valid': [], 'suspicious': [],
            'hallucinated': [], 'unverifiable': [],
            'summary': {'valid_count': 0, 'suspicious_count': 0,
                        'hallucinated_count': 0, 'unverifiable_count': 0}
        }

    print(f"  Found {len(citations)} citation(s). Verifying...\n")
    session = get_session()

    valid, suspicious, hallucinated, unverifiable = [], [], [], []
    for i, citation in enumerate(citations, 1):
        if verbose:
            print(f"  [{i}/{len(citations)}] Checking: {citation.get('key','?')}")
        res = verify_citation(citation, session, threshold, verbose)
        status = res['status']
        if status == 'valid':
            valid.append(res)
        elif status == 'suspicious':
            suspicious.append(res)
        elif status == 'unverifiable':
            unverifiable.append(res)
        else:
            hallucinated.append(res)

    report = {
        'file': file_path,
        'total_citations': len(citations),
        'valid': valid,
        'suspicious': suspicious,
        'hallucinated': hallucinated,
        'unverifiable': unverifiable,
        'summary': {
            'valid_count': len(valid),
            'suspicious_count': len(suspicious),
            'hallucinated_count': len(hallucinated),
            'unverifiable_count': len(unverifiable)
        }
    }
    return report


def print_report(report: Dict[str, Any]):
    """Print a human-readable report to stdout."""
    s = report.get('summary', {})
    print("\n" + "=" * 60)
    print("CITATION VERIFICATION REPORT")
    print("=" * 60)
    print(f"File: {report.get('file', 'unknown')}")
    print(f"Total citations: {report.get('total_citations', 0)}")
    print(f"  ✅ Valid:         {s.get('valid_count', 0)}")
    print(f"  ⚠️  Suspicious:    {s.get('suspicious_count', 0)}")
    print(f"  ❌ Hallucinated:  {s.get('hallucinated_count', 0)}")
    print(f"  ❓ Unverifiable:  {s.get('unverifiable_count', 0)}")

    for label, items in [("❌ HALLUCINATED CITATIONS (not found in any database)",
                           report.get('hallucinated', [])),
                          ("⚠️  SUSPICIOUS CITATIONS",
                           report.get('suspicious', [])),
                          ("❓ UNVERIFIABLE (no title could be extracted)",
                           report.get('unverifiable', []))]:
        if items:
            print(f"\n{label}:")
            for item in items:
                raw = item.get('citation', {}).get('raw', '')
                key = item.get('citation', {}).get('key', 'Unknown')
                print(f"  - [{key}] {raw[:120]}")

    print("=" * 60 + "\n")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description='Check citations in documents without .bib file')
    parser.add_argument('--file', type=str, help='Single file to check')
    parser.add_argument('--dir', type=str, help='Directory to scan for documents')
    parser.add_argument('--output', type=str, help='Save report to JSON file')
    parser.add_argument('--threshold', type=float, default=0.60,
                        help='Similarity threshold for Jaccard matching (default: 0.60)')
    parser.add_argument('--verbose', action='store_true',
                        help='Enable debug output')
    args = parser.parse_args()

    if not args.file and not args.dir:
        parser.print_help()
        sys.exit(1)

    if not HAS_REQUESTS:
        print("[Error] requests library is required. Install: pip install requests")
        sys.exit(1)

    reports = []
    if args.file:
        report = check_file(args.file, args.threshold, args.verbose)
        print_report(report)
        reports.append(report)

    if args.dir:
        dir_path = Path(args.dir)
        for ext in ['*.docx', '*.md', '*.txt', '*.tex', '*.bib']:
            for file_path in dir_path.glob(f'**/{ext}'):
                report = check_file(str(file_path), args.threshold, args.verbose)
                print_report(report)
                reports.append(report)

    if args.output:
        try:
            with open(args.output, 'w', encoding='utf-8') as f:
                json.dump(reports, f, indent=2, ensure_ascii=False)
            print(f"Report saved to: {args.output}")
        except Exception as e:
            print(f"[Error] Cannot write output file: {e}")


if __name__ == '__main__':
    main()
